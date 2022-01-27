from PySide2.QtWidgets import QApplication, QMessageBox, QFileDialog, QMainWindow, QPushButton, QPlainTextEdit, \
    QDateEdit
from PySide2.QtUiTools import QUiLoader
import os
import PySide2
import docx
from win32com import client
import json
import re
import xlrd
import jieba
import crawler

dirname = os.path.dirname(PySide2.__file__)
plugin_path = os.path.join(dirname, 'plugins', 'platforms')
os.environ['QT_QPA_PLATFORM_PLUGIN_PATH'] = plugin_path


class Stats:
    def __init__(self):
        self.ui = QUiLoader().load("ui/ui.ui")
        # 爬取按钮
        self.ui.findcase_button.clicked.connect(self.findcase)
        # 上传按钮
        self.ui.loadcase_button.clicked.connect(self.upload)
        # 清空按钮
        self.ui.clear_button.clicked.connect(self.cleartext)
        # 分析按钮
        self.ui.analyse_button.clicked.connect(self.analyse)
        # 保存按钮
        self.ui.save_button.clicked.connect(self.save)
        # 批量处理按钮
        self.ui.batchprocess_button.clicked.connect(self.batchprocess)


    def findcase(self):
        # 爬虫程序，start，end为date对象，对该时间内的案件进行爬取
        start=self.ui.startdate.date()
        end=self.ui.enddate.date()
        crawler.search(start, end)

    def upload(self):
        filePath, _ = QFileDialog.getOpenFileName(
            self.ui,  # 父窗口对象
            "选择你要上传的文件",  # 标题
            r"d:\\text",  # 起始目录
            "文本类型 (*.txt *.docx *.doc)"  # 选择类型过滤项，过滤内容在括号中
        )
        if filePath=="":
            pass
        else:
            if filePath.endswith(".docx"):
                pass
            else:
                # 把doc格式转换成docx格式
                self.doc2docx(filePath)
                os.remove(filePath)
                filePath = filePath[0:len(filePath) - 4] + ".docx"
            doc = docx.Document(filePath)
            content = ""
            for value in doc.paragraphs:  # 遍历文档的每一段
                content += value.text + "\n"  # 输出每一段的内容
            self.ui.casetext_edit.setPlainText(content)


    def doc2docx(self,fn):
        word = client.Dispatch("Word.Application")  # 打开word应用程序
        # for file in files:
        doc = word.Documents.Open(fn)  # 打开word文件
        doc.SaveAs("{}x".format(fn), 12)  # 另存为后缀为".docx"的文件，其中参数12或16指docx文件
        doc.Close()  # 关闭原来word文件
        word.Quit()

    def cleartext(self):
        self.ui.casetext_edit.clear()

    def save(self):
        # 将案件文本和标注分别保存为.txt和json文件
        label={}
        party=[]
        if self.ui.party_choice1.isChecked() :
            party.append(self.ui.party_choice1.text())
        if self.ui.party_choice2.isChecked() :
            party.append(self.ui.party_choice2.text())
        if self.ui.party_choice3.isChecked() :
            party.append(self.ui.party_choice3.text())
        if self.ui.party_choice4.isChecked() :
            party.append(self.ui.party_choice4.text())
        label["当事人"] =party
        label["当事人1性别"] = self.ui.gender_group1.checkedButton().text()
        label["当事人2性别"] = self.ui.gender_group2.checkedButton().text()
        label["当事人1民族"] = self.ui.nation_group1.checkedButton().text()
        label["当事人2民族"] = self.ui.nation_group2.checkedButton().text()
        label["当事人1出生地"] = self.ui.birthplace_group1.checkedButton().text()
        label["当事人2出生地"] = self.ui.birthplace_group2.checkedButton().text()
        cause=[]
        if self.ui.cause_choice1.isChecked() :
            cause.append(self.ui.cause_choice1.text())
        if self.ui.cause_choice1.isChecked() :
            cause.append(self.ui.cause_choice2.text())
        if self.ui.cause_choice1.isChecked() :
            cause.append(self.ui.cause_choice3.text())
        if self.ui.cause_choice1.isChecked() :
            cause.append(self.ui.cause_choice4.text())
        label["案由"]=cause
        court=[]
        if self.ui.court_choice1.isChecked() :
            court.append(self.ui.court_choice1.text())
        if self.ui.court_choice2.isChecked() :
            court.append(self.ui.court_choice2.text())
        if self.ui.court_choice3.isChecked() :
            court.append(self.ui.court_choice3.text())
        if self.ui.court_choice4.isChecked() :
            court.append(self.ui.court_choice4.text())
        label["相关法院"] = court
        jsonData = json.dumps(label)
        filePath = QFileDialog.getExistingDirectory(self.ui, "选择存储路径")
        if filePath=="":
            return
        fileObject = open(filePath+'/'+"".join(party)+"".join(cause)+'案标注'+'.json', 'w')
        fileObject.write(jsonData)
        fileObject.close()
        text_content=self.ui.casetext_edit.toPlainText()
        text_object=open(filePath+'/'+"".join(party)+"".join(cause)+'案文本'+'.txt', 'w')
        text_object.write(text_content)
        text_object.close()

    def batchprocess(self):
        # 对文件夹中的内容进行处理
        filePath = QFileDialog.getExistingDirectory(self.ui, "选择待处理的文件夹路径")
        # 正则表达式 ↓
        CriminalsObj = re.compile("被告人(?P<Criminals>.*?)[，,,][男,女]，.*?族")
        EthnicityObj = re.compile("[\u7537,\u5973][，,,](?P<Ethnicity>.*?)族.*?出生")
        BirthPlaceObj = re.compile("出生于(?P<BirthPlace>.*?)[，,,]")
        GenderObj = re.compile("被告人.*?[，,,](?P<Gender>[男,女]).*?族[，,,].*?出生")
        AccusationObj = re.compile("被告人[\u4e00-\u9fa5]{2,3}犯(?P<Accusation>.*?)一案[，,,]于.*?判决，认定")
        dic = {'Criminals': [], 'Ethnicity': [], 'BirthPlace': [], 'Gender': [], 'Courts': [], 'Accusation': []}

        files = os.listdir(filePath)
        # 取文书内容
        for file in files:
            title = file.split("/")[-1].replace(".docx", "").replace(".doc", "").replace(".txt", "")
            txt = docx.Document(filePath+"/"+file)
            paragraphs = txt.paragraphs
            file_content = ""
            for paragraph in paragraphs:
                file_content = file_content + paragraph.text

            dic['Criminals'] = CriminalsObj.findall(file_content)
            dic['Ethnicity'] = EthnicityObj.findall(file_content)
            dic['BirthPlace'] = BirthPlaceObj.findall(file_content)
            dic['Gender'] = GenderObj.findall(file_content)
            dic['Accusation'] = AccusationObj.findall(file_content)

            wb = xlrd.open_workbook('全国法院列表.xlsx')
            sh = wb.sheet_by_name('Sheet1')
            courts_list = sh.col_values(0)
            for court in courts_list:
                jieba.add_word(court)
            # 匹配相关法院
            courts = []
            seg = jieba.cut(file_content)
            for word in seg:
                if word.endswith("法院") and word != "法院" and word != "人民法院" and word != "高级人民法院" and word != "中级人民法院":
                    courts.append(word)
            courts = list(set(courts))  # 相关法院
            dic['Courts'] = courts

            json_str = json.dumps(dic)
            with open(filePath + "/" + title + "标注" + ".json", "wb") as json_file:
                json_file.write(json_str.encode())
            with open(filePath + "/" + title + ".txt", "wb") as txt_file:
                txt_file.write(file_content.encode())

    def analyse(self):
        # 对案件文本进行分析，更改tab页中的选项
        text=self.ui.casetext_edit.toPlainText()
        self.change_party(self.analyse_party(text))
        self.change_gender1(self.analyse_gender1(text))
        self.change_gender2(self.analyse_gender2(text))
        self.change_nation1(self.analyse_nation1(text))
        self.change_nation2(self.analyse_nation2(text))
        self.change_birthplace1(self.analyse_birthplace1(text))
        self.change_birthplace2(self.analyse_birthplace2(text))
        self.change_cause(self.analyse_cause(text))
        self.change_court(self.analyse_court(text))


    # 分析并更改当事人
    def change_party(self,choices):
        partychoice_list=[]
        partychoice_list.append(self.ui.party_choice1)
        partychoice_list.append(self.ui.party_choice2)
        partychoice_list.append(self.ui.party_choice3)
        partychoice_list.append(self.ui.party_choice4)
        for i in range(0,4):
            partychoice_list[i].setText(choices[i])

    def analyse_party(self,text):
        # 返回list 排序 大小为4
        CriminalsObj = re.compile("被告人(?P<Criminals>.*?)[，,,][男,女]，.*?族", re.S)
        Criminals = CriminalsObj.findall(text)
        while len(Criminals) < 4:
            Criminals.append("无")
        return Criminals


    # 分析并更改当事人1性别
    def change_gender1(self, choices):
        if choices[0]=="男" :
            self.ui.boy_choice1.setChecked(True)
            self.ui.girl_choice1.setChecked(False)
        else:
            self.ui.boy_choice1.setChecked(False)
            self.ui.girl_choice1.setChecked(True)

    def analyse_gender1(self, text):
        # 返回list，大小为1，比如list[0]=“男”
        GenderObj = re.compile("被告人.*?(?P<Gender>男|女).*?族", re.S)
        Gender = GenderObj.findall(text)
        return Gender[0:1]

    # 分析并更改当事人2性别
    def change_gender2(self, choices):
        if len(choices)==0:
            self.ui.boy_choice2.setCheckable(False)
            self.ui.girl_choice2.setCheckable(False)
            self.ui.boy_choice2.setText("无")
            self.ui.girl_choice2.setText("无")
            return
        if choices[0]=="男" :
            self.ui.boy_choice2.setChecked(True)
            self.ui.girl_choice2.setChecked(False)
        else:
            self.ui.boy_choice2.setChecked(False)
            self.ui.girl_choice2.setChecked(True)

    def analyse_gender2(self, text):
        # 返回list，大小为1，比如list[0]=“男” 若无返回空list
        GenderObj = re.compile("被告人.*?(?P<Gender>男|女).*?族", re.S)
        Gender = GenderObj.findall(text)
        if(len(Gender) > 1):   # 有多个当事人
            return Gender[1:2]
        else:
            return []
    # 分析并更改当事人1民族
    def change_nation1(self, choices):
        nationchoice_list=[]
        nationchoice_list.append(self.ui.nation_choice1)
        nationchoice_list.append(self.ui.nation_choice2)
        for i in range(0,2):
            nationchoice_list[i].setText(choices[i])

    def analyse_nation1(self, text):
        # 返回list 排序 大小为2
        EthnicityObj = re.compile("[\u7537,\u5973]，(?P<Ethnicity>.*?)族.*?出生", re.S)
        Ethnicity = EthnicityObj.findall(text)
        Ethnicity = [ethnicity + "族" for ethnicity in Ethnicity]
        ethnicity_list = []
        if len(Ethnicity) != 0:
            ethnicity_list.append(Ethnicity[0])
            ethnicity_list.append("无")
            return ethnicity_list
        else:
            return []

    # 分析并更改当事人2民族
    def change_nation2(self, choices):
        if len(choices)==0:
            self.ui.nation_choice12.setCheckable(False)
            self.ui.nation_choice22.setCheckable(False)
            self.ui.nation_choice12.setText("无")
            self.ui.nation_choice22.setText("无")
            return
        nationchoice_list=[]
        nationchoice_list.append(self.ui.nation_choice12)
        nationchoice_list.append(self.ui.nation_choice22)
        for i in range(0,2):
            nationchoice_list[i].setText(choices[i])

    def analyse_nation2(self, text):
        # 返回list 排序 大小为2 若无返回空list
        EthnicityObj = re.compile("[\u7537,\u5973]，(?P<Ethnicity>.*?)族.*?出生于", re.S)
        Ethnicity = EthnicityObj.findall(text)
        Ethnicity = [ethnicity + "族" for ethnicity in Ethnicity]
        ethnicity_list = []
        if len(Ethnicity) > 1:
            ethnicity_list.append(Ethnicity[1])
            ethnicity_list.append("无")
            return ethnicity_list
        else:
            return []

    # 分析并更改当事人1出生地
    def change_birthplace1(self, choices):
        birthplacechoice_list=[]
        birthplacechoice_list.append(self.ui.birthplace_choice1)
        birthplacechoice_list.append(self.ui.birthplace_choice2)
        for i in range(0,2):
            birthplacechoice_list[i].setText(choices[i])


    def analyse_birthplace1(self, text):
        # 返回list 排序 大小为2
        BirthPlaceObj = re.compile("出生于(?P<BirthPlace>.*?)，", re.S)
        BirthPlace = BirthPlaceObj.findall(text)
        birthplace_list = []
        if len(BirthPlace) != 0:
            birthplace_list.append(BirthPlace[0])
            birthplace_list.append("无")
            return birthplace_list
        else:
            return []

    # 分析并更改当事人2出生地
    def change_birthplace2(self, choices):
        if len(choices)==0:
            self.ui.birthplace_choice12.setCheckable(False)
            self.ui.birthplace_choice22.setCheckable(False)
            self.ui.birthplace_choice12.setText("无")
            self.ui.birthplace_choice22.setText("无")
            return
        birthplacechoice_list=[]
        birthplacechoice_list.append(self.ui.birthplace_choice12)
        birthplacechoice_list.append(self.ui.birthplace_choice22)
        for i in range(0,2):
            birthplacechoice_list[i].setText(choices[i])


    def analyse_birthplace2(self, text):
        # 返回list 排序 大小为2 若无返回空list
        BirthPlaceObj = re.compile("出生于(?P<BirthPlace>.*?)，", re.S)
        BirthPlace = BirthPlaceObj.findall(text)
        birthplace_list = []
        if len(BirthPlace) > 1:
            birthplace_list.append(BirthPlace[1])
            birthplace_list.append("无")
            return birthplace_list
        else:
            return []

    # 分析并更改案由
    def change_cause(self, choices):
        causechoice_list=[]
        causechoice_list.append(self.ui.cause_choice1)
        causechoice_list.append(self.ui.cause_choice2)
        causechoice_list.append(self.ui.cause_choice3)
        causechoice_list.append(self.ui.cause_choice4)
        for i in range(0,4):
            causechoice_list[i].setText(choices[i])

    def analyse_cause(self, text):
        # 返回list 排序 大小为4
        AccusationObj = re.compile("被告人[\u4e00-\u9fa5,\u3001]{2,10}犯(?P<Accusation>.*?)一案[，,,]于.*?判决，认定", re.S)
        Accusation = AccusationObj.findall(text)
        accusation_list = []
        try:
            accusation_list = Accusation[0].split("、")
            while len(accusation_list) < 4:
                accusation_list.append("无")
        except IndexError:
            while len(accusation_list) < 4:
                accusation_list.append("无")
        return [accusation.replace("罪", "") for accusation in accusation_list]

    # 分析并更改相关法院
    def change_court(self, choices):
        courtchoice_list=[]
        courtchoice_list.append(self.ui.court_choice1)
        courtchoice_list.append(self.ui.court_choice2)
        courtchoice_list.append(self.ui.court_choice3)
        courtchoice_list.append(self.ui.court_choice4)
        for i in range(0,4):
            courtchoice_list[i].setText(choices[i])

    def analyse_court(self, text):
        # 返回list 排序 大小为4
        # 将法院名称添加到jieba词典
        wb = xlrd.open_workbook('全国法院列表.xlsx')
        sh = wb.sheet_by_name('Sheet1')
        courts_list = sh.col_values(0)
        for court in courts_list:
            jieba.add_word(court)
        # 匹配相关法院
        courts = []
        seg = jieba.cut(text)
        for word in seg:
            if word.endswith("法院") and word != "法院" and word != "人民法院" and word != "高级人民法院" and word != "中级人民法院":
                courts.append(word)
        courts = list(set(courts))  # 相关法院
        while len(courts) < 4:
            courts.append("无")
        return courts

app = QApplication([])
stats = Stats()
stats.ui.show()
app.exec_()
